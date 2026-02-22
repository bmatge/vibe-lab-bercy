"""Convert Markdown text to a DOCX file (returns a BytesIO buffer)."""

import io
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Regex to split text into inline-formatted segments:
# **bold**, *italic*, `code`, or plain text
_INLINE_RE = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')


def _add_runs(paragraph, text: str):
    """Parse inline markdown and add formatted runs to a paragraph."""
    text = text.strip()
    pos = 0
    for m in _INLINE_RE.finditer(text):
        # Add plain text before this match
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])

        if m.group(2) is not None:  # **bold**
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3) is not None:  # *italic*
            run = paragraph.add_run(m.group(3))
            run.italic = True
        elif m.group(4) is not None:  # `code`
            run = paragraph.add_run(m.group(4))
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        pos = m.end()

    # Add remaining plain text
    if pos < len(text):
        paragraph.add_run(text[pos:])


def md_to_docx(md_text: str, title: str = '') -> io.BytesIO:
    doc = Document()

    # Basic style tweaks
    style = doc.styles['Normal']
    style.font.name = 'Marianne'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    if title:
        doc.add_heading(title, level=0)

    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            heading = doc.add_heading(level=level)
            _add_runs(heading, m.group(2))
            i += 1
            continue

        # Blockquote
        m = re.match(r'^>\s*(.*)', line)
        if m:
            # Collect consecutive blockquote lines
            bq_parts = [m.group(1)]
            i += 1
            while i < len(lines) and re.match(r'^>\s?(.*)', lines[i]):
                bq_parts.append(re.match(r'^>\s?(.*)', lines[i]).group(1))
                i += 1
            bq_text = ' '.join(p for p in bq_parts if p.strip())
            if bq_text.strip():
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                p.paragraph_format.left_indent = Cm(1)
                _add_runs(p, bq_text)
            continue

        # Horizontal rule
        if re.match(r'^---+\s*$', line):
            i += 1
            continue

        # Table: detect header row starting with |
        if line.strip().startswith('|') and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|', lines[i + 1]):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue

        # Bullet list
        m = re.match(r'^[-*]\s+(.*)', line)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            _add_runs(p, m.group(1))
            i += 1
            continue

        # Numbered list
        m = re.match(r'^\d+\.\s+(.*)', line)
        if m:
            p = doc.add_paragraph(style='List Number')
            _add_runs(p, m.group(1))
            i += 1
            continue

        # Blank line → skip
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_runs(p, line)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_table(doc: Document, table_lines: list[str]):
    """Parse markdown table lines and add a Word table."""
    def parse_row(line):
        cells = line.strip().strip('|').split('|')
        return [c.strip() for c in cells]

    rows = []
    for idx, line in enumerate(table_lines):
        if idx == 1 and re.match(r'^[\s\-:|]+$', line.strip().strip('|')):
            continue  # skip separator row
        rows.append(parse_row(line))

    if not rows:
        return

    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < n_cols:
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = ''  # clear default paragraph
                _add_runs(cell.paragraphs[0], cell_text)
