"""Tests de la conversion Markdown → DOCX."""
import io

from docx import Document

from server.md_to_docx import md_to_docx


def _read_docx(buf: io.BytesIO) -> Document:
    """Ouvre un buffer DOCX en tant que Document."""
    buf.seek(0)
    return Document(buf)


def test_basic_conversion():
    """Texte simple produit un DOCX valide."""
    buf = md_to_docx('Bonjour le monde', title='Test')
    doc = _read_docx(buf)
    texts = [p.text for p in doc.paragraphs]
    assert any('Bonjour le monde' in t for t in texts)


def test_title():
    """Le titre apparaît dans le document."""
    buf = md_to_docx('Contenu', title='Mon Titre')
    doc = _read_docx(buf)
    assert doc.paragraphs[0].text == 'Mon Titre'


def test_headings():
    """Les niveaux h1-h4 sont correctement convertis."""
    md = '# H1\n## H2\n### H3\n#### H4'
    buf = md_to_docx(md)
    doc = _read_docx(buf)
    heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert 'H1' in heading_texts
    assert 'H2' in heading_texts
    assert 'H3' in heading_texts
    assert 'H4' in heading_texts


def test_inline_formatting():
    """Bold, italic et code inline sont gérés."""
    md = '**bold** *italic* `code`'
    buf = md_to_docx(md)
    doc = _read_docx(buf)
    # Vérifie que le texte est présent
    full_text = ' '.join(p.text for p in doc.paragraphs)
    assert 'bold' in full_text
    assert 'italic' in full_text
    assert 'code' in full_text


def test_tables():
    """Un tableau markdown est converti en table Word."""
    md = '| A | B |\n|---|---|\n| 1 | 2 |\n| 3 | 4 |'
    buf = md_to_docx(md)
    doc = _read_docx(buf)
    assert len(doc.tables) >= 1
    table = doc.tables[0]
    assert table.rows[0].cells[0].text == 'A'
    assert table.rows[1].cells[1].text == '2'


def test_bullet_list():
    """Les listes à puces sont converties."""
    md = '- item un\n- item deux\n- item trois'
    buf = md_to_docx(md)
    doc = _read_docx(buf)
    list_items = [p.text for p in doc.paragraphs if 'List' in p.style.name]
    assert len(list_items) == 3


def test_numbered_list():
    """Les listes numérotées sont converties."""
    md = '1. premier\n2. deuxieme\n3. troisieme'
    buf = md_to_docx(md)
    doc = _read_docx(buf)
    list_items = [p.text for p in doc.paragraphs if 'Number' in p.style.name]
    assert len(list_items) == 3


def test_empty_input():
    """Un markdown vide produit un DOCX valide."""
    buf = md_to_docx('')
    doc = _read_docx(buf)
    assert doc is not None
